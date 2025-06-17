import os
import re
from datetime import datetime
from docx import Document

# 修改为你的 Vue3 项目路径
TARGET_DIR = "/Users/oncechen/IdeaProjects/ued_web_static_vue"

API_PATTERN = re.compile(r"""
    (?P<caller>axios(?:Instance)?)\.(?P<method>get|post)     # axios 或 axiosInstance
    \s*\(\s*['"](?P<url>/api/[^'"]+)['"]                     # /api 开头的路径
    (?:\s*,\s*(?P<params>[^)]+))?                            # 可选参数
    \s*\)
""", re.VERBOSE)

def extract_api_calls_with_comments(file_path):
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        lines = f.readlines()

    results = []
    for idx, line in enumerate(lines):
        match = API_PATTERN.search(line)
        if match and not match.group("url").endswith(".json"):
            comment = ""
            for offset in range(1, 4):
                if idx - offset >= 0:
                    prev_line = lines[idx - offset].strip()
                    if prev_line.startswith("//"):
                        comment = prev_line.lstrip("//").strip()
                        break
            results.append({
                "来源文件": os.path.relpath(file_path, TARGET_DIR),
                "请求方式": match.group("method").upper(),
                "API接口": match.group("url"),
                "传递参数": match.group("params").strip() if match.group("params") else "",
                "功能名称": comment
            })
    return results

def scan_directory(base_dir):
    seen = set()
    all_apis = []
    for root, _, files in os.walk(base_dir):
        for file in files:
            if file.endswith((".js", ".ts", ".vue")):
                path = os.path.join(root, file)
                for api in extract_api_calls_with_comments(path):
                    key = (api["API接口"], api["请求方式"])
                    if key not in seen:
                        seen.add(key)
                        all_apis.append(api)
    return all_apis

def export_to_word(api_data):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"API接口整理_{timestamp}.docx"
    doc = Document()
    doc.add_heading("Vue 项目 API 接口整理", level=1)

    for i, api in enumerate(api_data, start=1):
        doc.add_paragraph(f"{i}. 功能名称：{api['功能名称']}", style='List Number')
        doc.add_paragraph(f"   来源文件：{api['来源文件']}")
        doc.add_paragraph(f"   请求方式：{api['请求方式']}")
        doc.add_paragraph(f"   API 接口：{api['API接口']}")
        doc.add_paragraph(f"   传递参数：{api['传递参数']}")
        doc.add_paragraph("")

    doc.save(filename)
    print(f"✅ Word 文件已保存：{filename}")

if __name__ == "__main__":
    print(f"📁 正在扫描目录：{TARGET_DIR}")
    api_list = scan_directory(TARGET_DIR)
    if not api_list:
        print("❌ 没有找到任何 API 调用")
    else:
        for i, api in enumerate(api_list, start=1):
            print(f"#{i}")
            print(f"来源文件：{api['来源文件']}")
            print(f"请求方式：{api['请求方式']}")
            print(f"API接口：{api['API接口']}")
            print(f"传递参数：{api['传递参数']}")
            print(f"功能名称：{api['功能名称']}\n")

        confirm = input("是否导出为 Word 文件？（是 / 否）：").strip().lower()
        if confirm in ("是", "y", "yes"):
            export_to_word(api_list)